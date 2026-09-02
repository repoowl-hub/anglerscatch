from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading("Angler's Catch Code Walkthrough & Multiplayer Compatibility", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introduction
doc.add_heading("1. Overview", level=1)
doc.add_paragraph("This document provides a visual and conceptual breakdown of the Angler's Catch mod codebase. The mod introduces dynamic sizes for fish, visual scaling of entities, custom trophy displays, and proportional meat drops when butchering. Each section below explains the responsibilities of the core C# files in the project.")

# Architecture
doc.add_heading("2. File Breakdown", level=1)

# AnglersCatchModSystem.cs
doc.add_heading("AnglersCatchModSystem.cs (The Entry Point)", level=2)
p = doc.add_paragraph()
p.add_run("What it does: ").bold = True
p.add_run("This is the main entry point of the mod. It tells the game to register our custom classes (like FishEntityBehavior and ItemFishRaw) so that the Vintage Story engine can use them. It also loads our Harmony patches when the server starts.")
doc.add_paragraph("- Registers behaviors and item classes.")
doc.add_paragraph("- Initializes the Harmony patching system exclusively on the ServerSide.")

# HarmonyPatches.cs
doc.add_heading("HarmonyPatches.cs (The Butchering Logic)", level=2)
p = doc.add_paragraph()
p.add_run("What it does: ").bold = True
p.add_run("This file is responsible for intercepting the vanilla butchering mechanics. Because vanilla fish drop a fixed amount of meat, we use an \"Inventory Diff\" approach to give the player extra meat dynamically based on the fish's size.")
doc.add_paragraph("- Uses Harmony Prefix to count how much raw fish the player has before butchering finishes.")
doc.add_paragraph("- Calculates a filletMultiplier based on the size of the specific fish.")
doc.add_paragraph("- Uses Harmony Postfix to check how much meat the vanilla game gave the player, and then awards extra meat into the player's inventory to match the multiplier.")
doc.add_paragraph("- Thread Safety: Uses [ThreadStatic] variables. This is crucial for multiplayer servers so that if two players butcher fish at the exact same millisecond, their data does not conflict.")

# FishItemBehavior.cs
doc.add_heading("FishItemBehavior.cs (Items and Tooltips)", level=2)
p = doc.add_paragraph()
p.add_run("What it does: ").bold = True
p.add_run("This file handles how fish items behave in the inventory and the world. It provides the logic for dynamically translating and appending text to the item tooltips.")
doc.add_paragraph("- Calculates the size and weight of a fish.")
doc.add_paragraph("- Adds translated text (e.g., \"Trophy Fish\" or \"Meat Yield: +20%\") to the inventory tooltip using Lang.Get().")
doc.add_paragraph("- Calculates visual scaling factors based on the fish's percentile within its species limits.")

# FishSpeciesConfig.cs
doc.add_heading("FishSpeciesConfig.cs (Configuration Data)", level=2)
p = doc.add_paragraph()
p.add_run("What it does: ").bold = True
p.add_run("This acts as the master dictionary for all fish species. It holds minimum and maximum sizes, weight coefficients, and trophy thresholds.")
doc.add_paragraph("- Provides a GetRange() method that other files use to look up the stats of a given fish species.")

# ItemFish.cs
doc.add_heading("ItemFish.cs (Storage and Taxidermy Transitions)", level=2)
p = doc.add_paragraph()
p.add_run("What it does: ").bold = True
p.add_run("This ensures that fish keep their custom stats (size, species, caught-by UID) when transitioning between states, like being processed in a barrel.")
doc.add_paragraph("- When placed into a barrel, it saves the fish's attributes to the server's ObjectCache.")
doc.add_paragraph("- When the barrel processing is complete (e.g., turning into a Taxidermy fish), it restores those saved attributes onto the newly created taxidermy item.")

# Multiplayer Compatibility
doc.add_heading("3. Multiplayer Compatibility", level=1)
doc.add_paragraph("Yes, this mod is fully compatible with multiplayer servers. In fact, it was explicitly designed with server architecture in mind. Here is why:")
p1 = doc.add_paragraph(style='List Bullet')
p1.add_run("Thread Safety: ").bold = True
p1.add_run("As mentioned in HarmonyPatches.cs, the use of [ThreadStatic] ensures that concurrent actions by multiple players on a server are handled independently without race conditions.")
p2 = doc.add_paragraph(style='List Bullet')
p2.add_run("Server-Side Authority: ").bold = True
p2.add_run("Actions like generating random fish sizes or modifying barrel outputs are strictly guarded by world.Side == EnumAppSide.Server checks. This prevents clients from spoofing sizes or causing desyncs.")
p3 = doc.add_paragraph(style='List Bullet')
p3.add_run("Attribute Synchronization: ").bold = True
p3.add_run("The size and species data are stored in the Itemstack/Entity Attributes tree, which Vintage Story natively synchronizes between the server and all connected clients. This means visual scaling will display perfectly for everyone in the server.")

doc.save("AnglersCatch_Code_Walkthrough.docx")
print("Document saved.")
